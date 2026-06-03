"""
Document Processor — handles PDF, DOCX, TXT/MD file parsing.

Supports:
- PDF: PyMuPDF for text, pdfplumber for tables, VLM for image descriptions
- DOCX: python-docx + Unstructured
- TXT/MD: plain text reading
"""
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
import pandas as pd


@dataclass
class ProcessedDocument:
    """Normalized document unit after processing."""
    content: str
    doc_type: str          # "text" / "table" / "image_description"
    source: str            # original file path
    page: int              # 1-indexed page number
    position: int          # element position on page
    metadata: dict = field(default_factory=dict)


def process_document(file_path: str) -> list[ProcessedDocument]:
    """
    Route to the correct parser based on file extension.

    Returns a list of ProcessedDocument objects.
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return process_pdf(file_path)
    elif ext == ".docx":
        return process_docx(file_path)
    elif ext in (".txt", ".md"):
        return process_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def process_pdf(file_path: str) -> list[ProcessedDocument]:
    """
    Process a PDF file:
    1. PyMuPDF (fitz) extracts text per page
    2. pdfplumber extracts tables → Markdown
    3. VLM (Qwen-VL / GLM-4V) describes embedded images
    4. Content annotation with source, page, position metadata
    """
    import fitz  # PyMuPDF

    results: list[ProcessedDocument] = []
    doc = fitz.open(file_path)

    for page_num in range(len(doc)):
        page = doc[page_num]

        # A. Extract text with PyMuPDF
        text = page.get_text("text")
        if text.strip():
            results.append(ProcessedDocument(
                content=text.strip(),
                doc_type="text",
                source=file_path,
                page=page_num + 1,
                position=0,
                metadata={"extractor": "pymupdf"}
            ))

        # B. Extract tables with pdfplumber
        try:
            import pdfplumber
            # Re-open per-page to avoid stale objects
            with pdfplumber.open(file_path) as pdf:
                if page_num < len(pdf.pages):
                    pdf_page = pdf.pages[page_num]
                    tables = pdf_page.extract_tables()
                    for i, table in enumerate(tables):
                        if table and len(table) > 1:
                            header = table[0]
                            rows = table[1:]
                            df = pd.DataFrame(rows, columns=header)
                            md_table = dataframe_to_markdown(df)
                            results.append(ProcessedDocument(
                                content=md_table,
                                doc_type="table",
                                source=file_path,
                                page=page_num + 1,
                                position=i + 1,
                                metadata={"extractor": "pdfplumber", "table_index": i}
                            ))
        except Exception:
            pass  # pdfplumber is best-effort

        # C. Extract images with descriptions via VLM
        images = page.get_images(full=True)
        for i, img in enumerate(images):
            try:
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n > 4:  # CMYK → RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_bytes = pix.tobytes("png")
                description = describe_image_with_vlm(img_bytes)
                results.append(ProcessedDocument(
                    content=description,
                    doc_type="image_description",
                    source=file_path,
                    page=page_num + 1,
                    position=i + 1,
                    metadata={"extractor": "vlm", "image_index": i}
                ))
            except Exception:
                pass  # VLM is best-effort

    doc.close()
    return results


def process_docx(file_path: str) -> list[ProcessedDocument]:
    """Process a DOCX file using python-docx."""
    from docx import Document as DocxDocument

    results: list[ProcessedDocument] = []
    doc = DocxDocument(file_path)

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if full_text.strip():
        results.append(ProcessedDocument(
            content=full_text.strip(),
            doc_type="text",
            source=file_path,
            page=1,
            position=0,
            metadata={"extractor": "python-docx"}
        ))

    # Extract tables from docx
    for i, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows and len(rows) > 1:
            header = rows[0]
            data = rows[1:]
            df = pd.DataFrame(data, columns=header)
            md_table = dataframe_to_markdown(df)
            results.append(ProcessedDocument(
                content=md_table,
                doc_type="table",
                source=file_path,
                page=1,
                position=i + 1,
                metadata={"extractor": "python-docx", "table_index": i}
            ))

    return results


def process_text(file_path: str) -> list[ProcessedDocument]:
    """Process a plain text or markdown file."""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    if content.strip():
        return [ProcessedDocument(
            content=content.strip(),
            doc_type="text",
            source=file_path,
            page=1,
            position=0,
            metadata={"extractor": "builtin"}
        )]
    return []


def describe_image_with_vlm(image_bytes: bytes) -> str:
    """
    Describe an image using a vision-language model.

    Uses Qwen-VL or GLM-4V. Falls back gracefully if no API key.
    Prompt: "请详细描述这张图片的内容，包括图表、流程、架构等关键信息。"
    """
    try:
        from config import DASHSCOPE_API_KEY, OPENAI_API_KEY
        import base64

        encoded = base64.b64encode(image_bytes).decode("utf-8")
        data_url = f"data:image/png;base64,{encoded}"

        # Try Qwen-VL via DashScope
        if DASHSCOPE_API_KEY:
            try:
                from openai import OpenAI
                client = OpenAI(
                    api_key=DASHSCOPE_API_KEY,
                    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
                )
                resp = client.chat.completions.create(
                    model="qwen-vl-plus",
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "image_url", "image_url": {"url": data_url}},
                            {"type": "text", "text": "请详细描述这张图片的内容，包括图表、流程、架构等关键信息。"}
                        ]
                    }],
                    max_tokens=512
                )
                return resp.choices[0].message.content or ""
            except Exception:
                pass

        # Fallback: mark as unprocessed image
        return "[图片 - 未处理]"
    except Exception:
        return "[图片 - 描述失败]"


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    """Convert a pandas DataFrame to a Markdown table string."""
    return df.to_markdown(index=False)
